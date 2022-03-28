#modules imported
import pyttsx3

from docx import Document
from docx.shared import Inches

#function created for pyttx
def speak(text):
    pyttsx3.speak(text)


document = Document()

#add profile picture to document
document.add_picture(
    "profile-pic.png", width=Inches(2))

#add contact details to document
name = input('What is your name? ')
speak("Hello" + name + "how are you today? ")
speak("What is your number? ")
phone_num = input("What is your number? ")
email = input("What is your email address? ")

document.add_paragraph(
    name + " | " + phone_num + " | " + email)

#add about me to document
document.add_heading("About Me", level=1)
document.add_paragraph(input("Tell me about yourself? "))

#add work experiance to document
document.add_heading("Work Experiance", level=1)
p = document.add_paragraph()

company = input("Enter Company ")
from_date = input("From Date ")
end_date = input("To Date ")

p.add_run(company + " ").bold = True
p.add_run(from_date + " - " + end_date + "\n").italic = True

experiance_details = input(
    "Please enter detail of your work experiance at " + company + ": ")
p.add_run(experiance_details + "\n""\n")

#add more experiance to document
while True:
    has_more_experiances = input(
        "Do you have more experiances? Yes or No ")
    if has_more_experiances.lower() == "yes":
        company = input("Enter Company ")
        from_date = input("From Date ")
        end_date = input("To Date ")

        p.add_run(company + " ").bold = True
        p.add_run(from_date + " - " + end_date + "\n").italic = True

        experiance_details = input(
            "Please enter detail of your work experiance at " + company + ": ")
        p.add_run(experiance_details + "\n""\n")
    else:
        break


#add key skills to document
document.add_heading("Key Skills", level=1)
key_skills = (input("Please enter your key skills: "))
document.add_paragraph(key_skills,  style='List Bullet')

#add more skills to document
while True:
    has_more_skills = input(
        "Do you have more skills? Yes or No ")
    if has_more_skills.lower() == "yes":
        key_skills = (input("Enter key skill: "))
        document.add_paragraph(key_skills, style='List Bullet')
    else:
        break


document.save("cv.docx")